import streamlit as st
import openai
from docx import Document
from io import BytesIO
import pandas as pd

# --- Page Config ---
st.set_page_config(page_title="🧠 QP + Assessment Generator", layout="centered")

st.title("🧠 Smart Question Paper & Assessment Generator")

# --- API Key ---
api_key = st.text_input("Enter your OpenAI API Key", type="password")

# --- Subject Info ---
subject = st.text_input("Enter Subject Name")
available_cos = [f"CO{i+1}" for i in range(6)]
available_pos = [f"PO{i+1}" for i in range(12)]
available_blooms = ["Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create"]

num_units = st.slider("Number of Units", 1, 6, 3)

unit_data = []
for i in range(num_units):
    st.subheader(f"📘 Unit {i+1}")
    title = st.text_input(f"Unit {i+1} Title", key=f"title_{i}")
    bloom = st.selectbox(f"Bloom’s Level", available_blooms, key=f"bloom_{i}")
    marks = st.number_input(f"Marks for Unit {i+1}", min_value=1, max_value=20, value=5, key=f"marks_{i}")
    co = st.multiselect("Select COs", available_cos, key=f"co_{i}")
    po = st.multiselect("Select POs", available_pos, key=f"po_{i}")
    unit_data.append((title, bloom, marks, co, po))

# --- Question Generator ---
def generate_questions(subject, unit_title, bloom_level, marks):
    prompt = f"""
You are an AI question setter. Create questions for the subject "{subject}" from the unit "{unit_title}".
Bloom’s taxonomy level: {bloom_level}. Total marks: {marks}.
Generate 2 to 3 diverse questions (theory, applied, short).
"""
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=600
    )
    return response.choices[0].message.content.strip()

# --- Main Button ---
if st.button("🧠 Generate Everything"):
    if not api_key:
        st.error("Please enter your API key.")
    else:
        openai.api_key = api_key
        doc = Document()
        doc.add_heading("Smart Question Paper", 0)
        doc.add_paragraph(f"Subject: {subject}")

        blueprint = []
        doc.add_heading("Generated Questions", level=1)
        for i, (title, bloom, marks, co, po) in enumerate(unit_data, 1):
            doc.add_heading(f"Unit {i}: {title}", level=2)
            try:
                qns = generate_questions(subject, title, bloom, marks)
                doc.add_paragraph(qns)
                blueprint.append({
                    "Unit": f"Unit {i}",
                    "Bloom Level": bloom,
                    "Marks": marks,
                    "COs": ", ".join(co),
                    "POs": ", ".join(po)
                })
            except Exception as e:
                doc.add_paragraph(f"[Error generating questions: {e}]")

        # Add Blueprint Table
        if blueprint:
            doc.add_page_break()
            doc.add_heading("Assessment Blueprint", level=1)
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Unit'
            hdr_cells[1].text = 'Bloom Level'
            hdr_cells[2].text = 'Marks'
            hdr_cells[3].text = 'COs'
            hdr_cells[4].text = 'POs'

            for row in blueprint:
                cells = table.add_row().cells
                cells[0].text = row["Unit"]
                cells[1].text = row["Bloom Level"]
                cells[2].text = str(row["Marks"])
                cells[3].text = row["COs"]
                cells[4].text = row["POs"]

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Question Paper + Assessment Export Ready")
        st.download_button("📥 Download Final DOCX", data=buffer,
                           file_name="question_paper_and_blueprint.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")